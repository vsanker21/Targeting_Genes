#!/usr/bin/env python3
"""
Build a Word (.docx) report from GLIOMA-TARGET tier-1 results and pipeline summary.

Requires: pip install python-docx

  python scripts/export_glioma_target_results_docx.py [--top-n 80] [--output results/module7/glioma_target_results_report.docx]

Reads:
  - results/module7/glioma_target_tier1_welch.tsv
  - results/pipeline_results_index.json (summary only, optional)
"""

from __future__ import annotations

import argparse
import json
import os
import sys
from datetime import datetime, timezone
from pathlib import Path

import pandas as pd
import yaml

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:
    print("ERROR: install python-docx:  pip install python-docx", file=sys.stderr)
    raise SystemExit(2)


def repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


def data_root() -> Path:
    env = os.environ.get("GLIOMA_TARGET_DATA_ROOT", "").strip()
    if env:
        return Path(env)
    cfg = yaml.safe_load((repo_root() / "config" / "data_sources.yaml").read_text(encoding="utf-8"))
    return Path(cfg["data_root"].replace("/", os.sep))


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument(
        "--tier1-tsv",
        default="results/module7/glioma_target_tier1_welch.tsv",
        help="Repo-relative tier-1 TSV path.",
    )
    ap.add_argument(
        "--pipeline-index",
        default="results/pipeline_results_index.json",
        help="Repo-relative pipeline index JSON (optional).",
    )
    ap.add_argument("--top-n", type=int, default=80, help="Max rows to include as a table in the document.")
    ap.add_argument(
        "--output",
        default="results/module7/glioma_target_results_report.docx",
        help="Repo-relative output .docx path.",
    )
    args = ap.parse_args()
    rr = repo_root()
    t1_path = rr / args.tier1_tsv.replace("/", os.sep)
    if not t1_path.is_file():
        print(f"ERROR: missing {t1_path}", file=sys.stderr)
        return 1

    df = pd.read_csv(t1_path, sep="\t", low_memory=False)
    n_total = len(df)
    cols_show = [
        c
        for c in (
            "hgnc_symbol",
            "gene_id",
            "glioma_target_score",
            "glioma_target_tier",
            "gts_sub_E_norm",
            "gts_sub_M_norm",
            "gts_sub_D_norm",
            "gts_sub_N_norm",
            "delta_log2_expression",
            "depmap_crispr_median_gbm",
            "gts_evidence_tier",
            "outline_m22_known_gbm_driver",
        )
        if c in df.columns
    ]
    top_n = max(0, min(int(args.top_n), n_total))
    head = df[cols_show].head(top_n)

    idx_path = rr / args.pipeline_index.replace("/", os.sep)
    idx_summary: dict | None = None
    if idx_path.is_file():
        try:
            idx_summary = json.loads(idx_path.read_text(encoding="utf-8")).get("summary")
        except json.JSONDecodeError:
            idx_summary = None

    doc = Document()
    title = doc.add_heading("GLIOMA-TARGET results report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"Generated {datetime.now(tz=timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} · composite v1.1 (E + M + D + N)\n"
    )
    run.font.size = Pt(11)

    doc.add_heading("Summary", level=1)
    g_min = float(df["glioma_target_score"].min()) if "glioma_target_score" in df.columns else float("nan")
    g_max = float(df["glioma_target_score"].max()) if "glioma_target_score" in df.columns else float("nan")
    n_drv = int(df["outline_m22_known_gbm_driver"].fillna(False).astype(bool).sum()) if "outline_m22_known_gbm_driver" in df.columns else 0
    for line in (
        f"Tier-1 gene list (glioma_target_tier == 1): {n_total} genes.",
        f"glioma_target_score range: {g_min:.4f} – {g_max:.4f} (v1.1 cohort percentile ranks; see config/glioma_target_score.yaml).",
        f"Genes flagged as outline known GBM drivers in this tier-1 slice: {n_drv}.",
        f"Source table (repo-relative): {args.tier1_tsv}",
        f"Full ranked Welch table: results/module7/gts_candidate_table_welch_stub.tsv",
        f"Documentation: docs/GLIOMA_TARGET_SCORE.md",
        f"data_root (config/env): {data_root()}",
    ):
        doc.add_paragraph(line, style="List Bullet")

    if idx_summary:
        doc.add_heading("Pipeline index snapshot", level=1)
        doc.add_paragraph(
            f"Indexed paths: {idx_summary.get('n_paths_unique', '—')} unique; "
            f"existing: {idx_summary.get('n_existing', '—')}; "
            f"missing required: {idx_summary.get('n_missing_required', '—')}."
        )
        pdel = idx_summary.get("primary_deliverables") or []
        if pdel:
            doc.add_paragraph("Primary deliverables (from pipeline_inventory.yaml):", style="List Bullet")
            for item in pdel:
                if isinstance(item, dict):
                    doc.add_paragraph(
                        f"{item.get('id', '')}: {item.get('path_posix', '')} — {item.get('tier_filter', '')}",
                        style="List Bullet 2",
                    )

    doc.add_heading(f"Top {top_n} tier-1 genes (by table order)", level=1)
    doc.add_paragraph(
        "Columns are a subset for readability; open the source TSV for all columns and all genes."
    )

    table = doc.add_table(rows=1, cols=len(cols_show))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(cols_show):
        hdr[i].text = col
    for _, row in head.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols_show):
            v = row[col]
            if pd.isna(v):
                cells[i].text = ""
            elif isinstance(v, bool):
                cells[i].text = "True" if v else "False"
            else:
                cells[i].text = str(round(float(v), 6)) if isinstance(v, float) else str(v)

    doc.add_heading("Notes", level=1)
    doc.add_paragraph(
        "gts_evidence_tier is a separate screen (DE + DepMap threshold) from glioma_target_tier "
        "(composite score bands). Tier-1 here is defined by glioma_target_tier == 1 only."
    )
    doc.add_paragraph(
        "This document is reproducible: re-run after refreshing outputs with "
        "python scripts/export_gts_candidate_table.py && python scripts/export_glioma_target_results_docx.py"
    )

    out_path = rr / args.output.replace("/", os.sep)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
